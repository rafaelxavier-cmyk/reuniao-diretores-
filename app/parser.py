"""
parser.py – Extrai a estrutura de uma pauta (DOCX ou PDF) para um dicionário
             que o generator.py transforma em slides.

Estrutura esperada na pauta:
  - Cabeçalho: "PAUTA – REUNIÃO DE DIRETORES", Data, Objetivo
  - Seção de agenda: lista numerada com "X. Nome – YY min"
  - Seções de conteúdo: "1. Título" com bullets e sub-itens
"""

import re
from pathlib import Path


# ── DOCX ──────────────────────────────────────────────────────────────────
def parse_docx(path: str) -> dict:
    from docx import Document
    doc = Document(path)

    pauta = {
        "titulo":   "REUNIÃO DE DIRETORES",
        "data":     "",
        "objetivo": "",
        "agenda":   [],   # [{"num": 1, "nome": "Cultura", "tempo": "15 min"}]
        "secoes":   [],   # [{"num": 1, "titulo": "Cultura", "tempo": "15 min",
                          #   "subtitulo": "", "items": ["...", ...],
                          #   "subsecoes": [{"titulo": "", "items": []}]}]
    }

    current_section = None
    current_sub     = None
    in_agenda_block = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""
        is_heading  = "heading" in style_name
        heading_lvl = 0
        if is_heading:
            m = re.search(r"heading\s+(\d+)", style_name)
            heading_lvl = int(m.group(1)) if m else 1

        # ── Cabeçalho da pauta ─────────────────────────────────────────────
        if re.match(r"pauta\s*[–-]?\s*reuni", text, re.I):
            pauta["titulo"] = text
            continue

        m = re.match(r"data\s*:\s*(.+)", text, re.I)
        if m:
            pauta["data"] = m.group(1).strip()
            continue

        m = re.match(r"objetivo\s*:\s*(.+)", text, re.I)
        if m:
            pauta["objetivo"] = m.group(1).strip()
            continue

        # ── Bloco de agenda ────────────────────────────────────────────────
        if re.match(r"^pauta\s*$", text, re.I) or re.match(r"^agenda\s*$", text, re.I):
            in_agenda_block = True
            continue

        agenda_match = re.match(
            r"^(\d+)[.\)]\s*(.+?)(?:\s*[-–]\s*(\d+\s*min))?$", text, re.I
        )
        if in_agenda_block and agenda_match:
            num  = int(agenda_match.group(1))
            nome = agenda_match.group(2).strip()
            tempo = agenda_match.group(3) or ""
            pauta["agenda"].append({"num": num, "nome": nome, "tempo": tempo})
            continue

        # Intervalo na agenda
        if in_agenda_block and re.match(r"intervalo", text, re.I):
            m_t = re.search(r"(\d+\s*min)", text, re.I)
            pauta["agenda"].append({
                "num": 0, "nome": "Intervalo",
                "tempo": m_t.group(1) if m_t else "20 min"
            })
            continue

        # ── Títulos de seção (H1 / pattern "N. Título") ───────────────────
        sec_match = re.match(r"^(\d+)[.\)]\s*(.+)", text)
        if sec_match and (heading_lvl == 1 or (heading_lvl == 0 and len(text) < 60)):
            num   = int(sec_match.group(1))
            title = sec_match.group(2).strip()
            # Pular se já está na agenda e é só repetição
            tempo = _find_tempo_for(num, pauta["agenda"])
            in_agenda_block = False
            current_section = {
                "num":       num,
                "titulo":    title,
                "tempo":     tempo,
                "subtitulo": "",
                "items":     [],
                "subsecoes": [],
            }
            current_sub = None
            pauta["secoes"].append(current_section)
            continue

        # ── Sub-seção (H2 / H3) ───────────────────────────────────────────
        if current_section and (heading_lvl in (2, 3) or
                                 (heading_lvl == 0 and _is_subsection_candidate(para))):
            current_sub = {"titulo": text, "items": []}
            current_section["subsecoes"].append(current_sub)
            if not current_section["subtitulo"]:
                current_section["subtitulo"] = text
            continue

        # ── Bullets / listas ──────────────────────────────────────────────
        is_list = (para.style and "list" in para.style.name.lower()) or \
                  text.startswith(("•", "-", "●", "◦", "▸", "–"))

        if is_list or (current_section and not is_heading and not sec_match):
            clean = re.sub(r"^[•\-●◦▸–]\s*", "", text)
            if current_sub:
                current_sub["items"].append(clean)
            elif current_section:
                current_section["items"].append(clean)

    # Se não encontrou seções via heading, tenta extrair de parágrafos em negrito
    if not pauta["secoes"]:
        pauta["secoes"] = _fallback_bold_sections(doc, pauta["agenda"])

    return pauta


# ── PDF ───────────────────────────────────────────────────────────────────
def parse_pdf(path: str) -> dict:
    import pdfplumber

    raw_lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                raw_lines.extend(text.split("\n"))

    # Reusa a lógica de texto simples
    return _parse_lines(raw_lines)


# ── INTERNAL ──────────────────────────────────────────────────────────────
def _parse_lines(lines: list) -> dict:
    pauta = {
        "titulo":   "REUNIÃO DE DIRETORES",
        "data":     "",
        "objetivo": "",
        "agenda":   [],
        "secoes":   [],
    }

    current_section = None
    current_sub     = None
    last_sub_num    = 0   # último número de sub-item visto na seção atual
    # Fases: 'header' → 'agenda' → 'content'
    phase = "header"

    for raw in lines:
        text = raw.strip()
        if not text:
            continue

        # ── Cabeçalho ────────────────────────────────────────────────────
        if re.match(r"pauta\s*[–\-]?\s*reuni", text, re.I):
            pauta["titulo"] = text
            continue

        m = re.match(r"data\s*:\s*(.+)", text, re.I)
        if m:
            pauta["data"] = m.group(1).strip()
            continue

        m = re.match(r"objetivo\s*:\s*(.+)", text, re.I)
        if m:
            pauta["objetivo"] = m.group(1).strip()
            continue

        # ── Transição para bloco de agenda ────────────────────────────────
        if re.match(r"^pauta\s*$", text, re.I):
            phase = "agenda"
            continue

        # ── Bloco de AGENDA: itens com tempo (X. Nome - Y min) ───────────
        # Só captura enquanto estiver na fase "agenda" E o item tiver " min"
        if phase == "agenda":
            # Intervalo dentro da agenda
            if re.match(r"intervalo", text, re.I):
                m_t = re.search(r"(\d+\s*min)", text, re.I)
                pauta["agenda"].append({
                    "num": 0, "nome": "Intervalo",
                    "tempo": m_t.group(1) if m_t else "20 min"
                })
                continue

            # Linha com tempo explícito → pertence à agenda
            ag_m = re.match(r"^(\d+)[.\)]\s*(.+?)\s*[-–]\s*(\d+\s*min)\s*$", text, re.I)
            if ag_m:
                pauta["agenda"].append({
                    "num":   int(ag_m.group(1)),
                    "nome":  ag_m.group(2).strip().rstrip("-–").strip(),
                    "tempo": ag_m.group(3),
                })
                continue

            # Linha numerada SEM tempo → fim da agenda, começa conteúdo
            sec_m = re.match(r"^(\d+)[.\)]\s+(\S.{1,60})$", text)
            if sec_m:
                phase = "content"
                # Cai para o bloco de conteúdo abaixo (sem continue)
            else:
                # Linha de horário, FIM etc. — pula
                continue

        # ── Bloco de CONTEÚDO ─────────────────────────────────────────────
        if phase == "content":
            # Seção principal: "N. Título" (linha curta, começa com número)
            sec_m = re.match(r"^(\d+)[.\)]\s*(\S.{1,70})$", text)
            if sec_m:
                num   = int(sec_m.group(1))
                title = sec_m.group(2).strip()
                # Decide se é seção top-level ou sub-item:
                # Regra 1 – sequencial: se num == last_sub_num+1, é sub-item (continuação de lista)
                # Regra 2 – agenda: se (nome bate OU número ainda não criado) E não é sequencial
                created_nums  = {s["num"] for s in pauta["secoes"]}
                agenda_nums   = {a["num"] for a in pauta["agenda"] if a["num"] > 0}
                name_match    = _matches_agenda(num, title, pauta["agenda"])
                num_match     = (num in agenda_nums) and (num not in created_nums)
                # Sequencial só bloca números pequenos (≤4): sub-listas raramente
                # passam de 4 itens, mas seções de agenda frequentemente chegam a 5-6
                is_sequential = (current_section
                                 and (num == last_sub_num + 1)
                                 and num <= 4)
                # name_match sempre cria seção; num_match cria se não for sequencial
                is_top_level  = name_match or (num_match and not is_sequential)
                if is_top_level:
                    last_sub_num = 0  # reset ao entrar em nova seção
                    tempo = _find_tempo_for(num, pauta["agenda"])
                    current_section = {
                        "num":       num,
                        "titulo":    title,
                        "tempo":     tempo,
                        "subtitulo": "",
                        "items":     [],
                        "subsecoes": [],
                    }
                    current_sub = None
                    pauta["secoes"].append(current_section)
                    continue
                else:
                    # Sub-item numerado → trata como bullet
                    last_sub_num = num  # atualiza o contador sequencial
                    clean = re.sub(r"^\d+[.\)]\s*", "", text)
                    if current_sub:
                        current_sub["items"].append(clean)
                    elif current_section:
                        current_section["items"].append(clean)
                    continue

            if not current_section:
                continue

            # Bullet explícito
            if text.startswith(("●", "•", "◦", "▸", "→")):
                clean = re.sub(r"^[●•◦▸→]\s*", "", text)
                if current_sub:
                    current_sub["items"].append(clean)
                else:
                    current_section["items"].append(clean)
                continue

            # Sub-seção: linha curta, não começa com bullet, sem ":" no meio
            # e não parece ser apenas corpo de texto longo
            is_short_title = (
                len(text) < 70
                and not text[0].isdigit()
                and not text.startswith(("http", "https"))
                and re.match(r"^[A-Za-zÀ-ž\U00000080-\U0001FFFF]", text)
                and text[-1] not in ".,"
            )
            if is_short_title and not current_section["subtitulo"]:
                current_section["subtitulo"] = text
                current_sub = {"titulo": text, "items": []}
                current_section["subsecoes"].append(current_sub)
                continue

            # Corpo de texto / linha genérica
            if current_sub:
                current_sub["items"].append(text)
            else:
                current_section["items"].append(text)

    return pauta


def _matches_agenda(num: int, title: str, agenda: list) -> bool:
    """Retorna True se (num, title) corresponde a um item da agenda principal."""
    title_lower = _normalize(title)
    for item in agenda:
        if item["num"] != num:
            continue
        nome_lower = _normalize(item["nome"])
        # Aceita se há sobreposição de palavras significativas (≥1 palavra com 4+ chars)
        words_title = set(w for w in title_lower.split() if len(w) >= 4)
        words_nome  = set(w for w in nome_lower.split()  if len(w) >= 4)
        if words_title & words_nome:
            return True
        # Fallback: um começa com o outro
        if title_lower.startswith(nome_lower[:8]) or nome_lower.startswith(title_lower[:8]):
            return True
    return False


def _normalize(s: str) -> str:
    """Remove acentos e converte para minúsculas para comparação."""
    import unicodedata
    nfkd = unicodedata.normalize("NFKD", s)
    return "".join(c for c in nfkd if not unicodedata.combining(c)).lower()


def _find_tempo_for(num: int, agenda: list) -> str:
    for item in agenda:
        if item["num"] == num:
            return item["tempo"]
    return ""


def _is_subsection_candidate(para) -> bool:
    """Retorna True se o parágrafo parece um sub-título (negrito, curto)."""
    text = para.text.strip()
    if not text or len(text) > 80:
        return False
    for run in para.runs:
        if run.bold:
            return True
    return False


def _fallback_bold_sections(doc, agenda: list) -> list:
    """Extrai seções a partir de parágrafos em negrito quando não há headings."""
    secoes = []
    current = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_bold = any(r.bold for r in para.runs if r.text.strip())
        sec_m = re.match(r"^(\d+)[.\)]\s+(.+)", text)
        if is_bold and sec_m:
            num = int(sec_m.group(1))
            current = {
                "num":       num,
                "titulo":    sec_m.group(2).strip(),
                "tempo":     _find_tempo_for(num, agenda),
                "subtitulo": "",
                "items":     [],
                "subsecoes": [],
            }
            secoes.append(current)
        elif current:
            clean = re.sub(r"^[●•\-◦▸]\s*", "", text)
            current["items"].append(clean)
    return secoes


# ── API pública ───────────────────────────────────────────────────────────
def parse(path: str) -> dict:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return parse_docx(path)
    elif ext == ".pdf":
        return parse_pdf(path)
    else:
        raise ValueError(f"Formato não suportado: {ext}. Use .docx ou .pdf")
